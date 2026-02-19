from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from typing import Optional
import PyPDF2
import io
import os
import json
import re
import tempfile
import subprocess
import shutil
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from groq import Groq

# Optional imports
try:
    import requests
    from bs4 import BeautifulSoup
    SCRAPING_AVAILABLE = True
except ImportError:
    SCRAPING_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

load_dotenv()
groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "https://resume-optimizer.vercel.app", "https://resume-optimizer-azure.vercel.app"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def read_root():
    return {"message": "Hello from Resume Optimizer Backend!"}

@app.get("/health")
def health_check():
    return {
        "status": "healthy",
        "ai_provider": "groq",
        "available_models": len(GROQ_MODELS),
        "primary_model": GROQ_MODELS[0]
    }

# ==================== GROQ MULTI-MODEL SUPPORT ====================

GROQ_MODELS = [
    "llama-3.3-70b-versatile",   # Best quality — primary (128k context)
    "llama-3.1-8b-instant",      # Fast fallback (128k context)
    "gemma2-9b-it",              # Google fallback (8k context)
]

# ── Round-robin counter so consecutive calls use different models
# This avoids hammering llama-3.3-70b on every single call
_model_counter = 0

def analyze_with_groq(prompt: str, max_tokens: int = 2000, prefer_large: bool = True) -> str:
    """
    Call Groq API with automatic model fallback.
    prefer_large=True  → start from llama-3.3-70b (best quality, use for JSON tasks)
    prefer_large=False → round-robin across all models (use for text tasks)
    """
    global _model_counter

    if prefer_large:
        # Always try best model first for JSON-heavy tasks
        ordered = GROQ_MODELS[:]
    else:
        # Rotate starting model to spread load
        start = _model_counter % len(GROQ_MODELS)
        ordered = GROQ_MODELS[start:] + GROQ_MODELS[:start]
        _model_counter += 1

    last_error = None
    for model in ordered:
        try:
            print(f"🤖 Trying Groq: {model}")
            chat_completion = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=model,
                temperature=0.7,
                max_tokens=max_tokens,
            )
            response = chat_completion.choices[0].message.content
            print(f"✅ Success with {model}")
            return response
        except Exception as e:
            error_msg = str(e).lower()
            last_error = str(e)
            print(f"⚠️ {model} failed: {error_msg[:120]}")
            continue

    raise Exception(
        f"All Groq models exhausted. Last error: {last_error}\n"
        f"Wait a few minutes — Groq free tier resets every minute per model."
    )

# ==================== ROBUST JSON EXTRACTION ====================

def extract_json_from_response(response: str) -> dict:
    """
    Robustly extract a JSON object from an LLM response.
    Handles:
      - Markdown code fences (```json ... ```)
      - Extra text before/after the JSON object
      - Truncated JSON (tries to close open braces)
      - 'Extra data' errors from llama-3.1 outputting text after closing brace
    """
    if not response:
        return None

    text = response.strip()

    # 1. Strip markdown code fences
    text = re.sub(r'^```(?:json)?\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\s*```$', '', text, flags=re.MULTILINE)
    text = text.strip()

    # 2. Extract the FIRST complete JSON object only
    #    Find the first { and match to its closing }
    brace_start = text.find('{')
    if brace_start == -1:
        return None

    depth = 0
    in_string = False
    escape_next = False
    end_pos = -1

    for i, ch in enumerate(text[brace_start:], start=brace_start):
        if escape_next:
            escape_next = False
            continue
        if ch == '\\' and in_string:
            escape_next = True
            continue
        if ch == '"' and not escape_next:
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                end_pos = i + 1
                break   # Stop at FIRST complete object — ignore any trailing text

    if end_pos == -1:
        # JSON was truncated — try to close open braces
        print("⚠️ JSON appears truncated, attempting repair...")
        partial = text[brace_start:]
        # Count unclosed braces
        open_braces = partial.count('{') - partial.count('}')
        repaired = partial + (']' if partial.rstrip().endswith('"') else '') + ('}' * open_braces)
        try:
            return json.loads(repaired)
        except Exception:
            return None

    json_str = text[brace_start:end_pos]

    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        print(f"⚠️ JSON parse error: {e}")
        # Last resort: try json.loads with strict=False won't help,
        # but try removing trailing commas (common LLM mistake)
        cleaned = re.sub(r',\s*([}\]])', r'\1', json_str)
        try:
            return json.loads(cleaned)
        except Exception:
            return None

# ==================== DOCUMENT EXTRACTION ====================

def extract_text_from_pdf_robust(file_content: bytes) -> str:
    if PDFPLUMBER_AVAILABLE:
        try:
            import pdfplumber
            pdf_file = io.BytesIO(file_content)
            text_parts = []
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                row_text = [str(cell).strip() for cell in row if cell]
                                if row_text:
                                    text_parts.append(' | '.join(row_text))
            result = '\n'.join(text_parts).strip()
            if len(result) > 100:
                print(f"✅ PDF extracted: {len(result)} chars")
                return result
        except Exception as e:
            print(f"pdfplumber failed: {e}")

    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = [page.extract_text() for page in pdf_reader.pages if page.extract_text()]
        result = '\n'.join(text).strip()
        print(f"✅ PDF extracted: {len(result)} chars")
        return result
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"

def extract_text_from_docx_robust(file_content: bytes) -> str:
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        full_text = []

        for section in doc.sections:
            for para in section.header.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

        for element in doc.element.body:
            if element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._element == element and para.text.strip():
                        full_text.append(para.text.strip())
                        break
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                full_text.append(' | '.join(row_text))
                        full_text.append('')
                        break

        for section in doc.sections:
            for para in section.footer.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

        result = '\n'.join(full_text).strip()
        print(f"✅ DOCX extracted: {len(result)} chars")
        return result
    except Exception as e:
        print(f"Advanced extraction failed: {e}")
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        text.append(' | '.join(row_data))
            return '\n'.join(text)
        except:
            return f"Error extracting DOCX: {str(e)}"

def extract_text_from_file(filename: str, file_content: bytes) -> str:
    filename_lower = filename.lower()
    print(f"📄 Extracting: {filename}")
    try:
        if filename_lower.endswith('.pdf'):
            return extract_text_from_pdf_robust(file_content)
        elif filename_lower.endswith('.docx'):
            return extract_text_from_docx_robust(file_content)
        elif filename_lower.endswith('.txt'):
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except:
                    continue
            return file_content.decode('utf-8', errors='ignore')
        else:
            return f"Unsupported format: {filename}"
    except Exception as e:
        return f"Error: {str(e)}"

# ==================== AI ANALYSIS FUNCTIONS ====================

def analyze_job_description(jd_content: str) -> dict:
    prompt = f"""Analyze this job description concisely.

Job Description:
{jd_content}

Provide:
1. **Top 10 Required Skills**
2. **Years of Experience**
3. **Must-Have Qualifications**
4. **Key Responsibilities** (top 5)
5. **ATS Keywords** (top 15)

Be concise and clear."""
    # Text task — rotate models to spread load
    analysis = analyze_with_groq(prompt, max_tokens=2000, prefer_large=False)
    return {"analysis": analysis, "success": True}

def run_full_analysis(jd_content: str, resume_content: str) -> dict:
    """Single API call covering ATS score, gaps, rewrites, skills, summary, and interview prep."""
    prompt = f"""You are a senior technical recruiter and expert resume coach. Analyze this resume against this job description thoroughly.

JOB DESCRIPTION:
{jd_content[:2500]}

RESUME:
{resume_content[:3000]}

Return ALL of the following sections. Use the EXACT section headers shown. Never skip a section.

###ATS_SCORE###
Overall ATS Score: [write a number]/100
Breakdown: Keyword Match ([number]/40), Experience ([number]/30), Skills ([number]/20), Education ([number]/10)
Top 5 Missing Keywords: [list them]
Top 3 Strengths: [list them]
Top 3 Improvements: [list them]

###CRITICAL_GAPS###
List 3-5 specific gaps — skills, experience, or keywords the JD requires that this resume lacks or undershows. Be concrete, not generic.

###REWRITES###
Pick the 3-5 weakest or vaguest bullet points from the resume. Rewrite each.
Use EXACTLY this format for each:
Original: <copy the exact bullet text from the resume, word for word>
Improved: <stronger version: action verb + specific metric + clear impact>

###SKILLS###
List skill/tool names from the JD missing from the resume (plausible for this candidate only).
One per line, name only, no explanation:
- SkillName

###SUMMARY###
Write a 2-3 sentence professional summary tailored to this exact JD and this candidate's background.

###INTERVIEW_PREP###
Technical Questions (5 specific to this JD):
Behavioral Questions (3):
Questions to Ask the Interviewer (3):
Topics to Study Before Interview (3):

Be specific to THIS resume and THIS job. Never fabricate experience."""

    response = analyze_with_groq(prompt, max_tokens=4000, prefer_large=False)
    return {"full_analysis": response, "success": True}

def generate_youtube_resources(jd_content: str) -> list:
    prompt = f"""List 5 technical topics to study for this job.

Job: {jd_content[:1000]}

List ONLY topic names, one per line. No numbering, no extra text."""
    topics_text = analyze_with_groq(prompt, max_tokens=500, prefer_large=False)
    topics = [t.strip().lstrip('0123456789.-) ') for t in topics_text.split('\n') if t.strip()][:5]
    return [{
        "topic": topic,
        "url": f"https://www.youtube.com/results?search_query={topic.replace(' ', '+')}+tutorial"
    } for topic in topics]

# ==================== RESUME PARSING ====================

def parse_resume_to_structured_data(resume_text: str) -> dict:
    prompt = f"""Extract resume information as JSON.

Resume:
{resume_text}

Return ONLY a valid JSON object with NO extra text before or after it:
{{
  "name": "Full Name",
  "phone": "Phone",
  "email": "Email",
  "location": "City, State",
  "linkedin": "URL or empty string",
  "github": "URL or empty string",
  "summary": "Professional summary or empty string",
  "education": [{{"degree": "BE", "institution": "University", "year": "2020"}}],
  "technical_skills": {{
    "languages": ["Python"],
    "databases": ["SQL"],
    "frameworks": [],
    "cloud": [],
    "tools": []
  }},
  "work_experience": [{{
    "title": "Title",
    "company": "Company",
    "location": "City",
    "duration": "2020-Present",
    "achievements": ["Achievement 1"]
  }}],
  "projects": [{{"name": "Project", "description": "Desc", "technologies": "Tech"}}],
  "certifications": [],
  "achievements": []
}}

IMPORTANT: Return ONLY the JSON object. No explanation, no markdown, no text outside the braces."""

    try:
        # Use large model for JSON tasks — accuracy matters here
        response = analyze_with_groq(prompt, max_tokens=4000, prefer_large=True)
        parsed_data = extract_json_from_response(response)
        if parsed_data:
            print(f"✅ Parsed: {parsed_data.get('name', 'Unknown')}")
            return parsed_data
        else:
            print("❌ Could not extract valid JSON from parse response")
            return None
    except Exception as e:
        print(f"❌ Parse error: {str(e)}")
        return None

def enhance_resume_data_with_ai(resume_data: dict, ai_suggestions: str, jd_analysis: str) -> dict:
    prompt = f"""You are a professional resume writer and grammar expert. Enhance the resume JSON.

ORIGINAL RESUME JSON:
{json.dumps(resume_data, indent=2)[:3000]}

SUGGESTIONS TO APPLY:
{ai_suggestions[:1000]}

STRICT RULES — follow every one:
1. PRESERVE all fields exactly: name, email, phone, location, linkedin, github, education, certifications, achievements, projects, technical_skills
2. NEVER remove or empty any field that has data in the original — especially "achievements" and "certifications"
3. Fix ALL grammar, spelling, and punctuation errors in every text field
4. Rewrite work experience bullet points in STAR format with action verbs and metrics where possible
5. Do NOT invent skills, jobs, or experience not present in the original
6. Keep the same JSON structure — do not add or remove keys

Return ONLY a valid JSON object with ALL original fields preserved. No markdown, no explanation, nothing outside the braces."""

    try:
        # Large model for JSON — accuracy critical
        response = analyze_with_groq(prompt, max_tokens=4000, prefer_large=True)
        enhanced_data = extract_json_from_response(response)

        if not enhanced_data:
            print("⚠️ Enhancement returned invalid JSON — using original")
            return {
                "original": resume_data,
                "enhanced": resume_data,
                "changes": ["Enhancement failed — using original data"],
                "success": False
            }

        # Safety net: restore any fields the AI silently dropped
        critical_fields = ['achievements', 'certifications', 'education', 'projects',
                           'name', 'email', 'phone', 'location', 'linkedin', 'github']
        for field in critical_fields:
            orig_val = resume_data.get(field)
            enh_val  = enhanced_data.get(field)
            if orig_val and (not enh_val or enh_val == [] or enh_val == ''):
                enhanced_data[field] = orig_val
                print(f"🔧 Restored dropped field: {field}")

        # Track changes
        changes = ["Fixed grammar and spelling across all sections"]
        orig_skills = set()
        new_skills = set()
        for category in ['languages', 'databases', 'frameworks', 'cloud', 'tools']:
            orig_skills.update(resume_data.get('technical_skills', {}).get(category, []))
            new_skills.update(enhanced_data.get('technical_skills', {}).get(category, []))

        added_skills = new_skills - orig_skills
        if added_skills:
            changes.append(f"Added {len(added_skills)} skills: {', '.join(list(added_skills)[:3])}")

        orig_exp = resume_data.get('work_experience', [])
        new_exp = enhanced_data.get('work_experience', [])
        if orig_exp and new_exp:
            changes.append("Enhanced work experience bullets with STAR format")

        print("✅ Enhanced successfully")
        return {
            "original": resume_data,
            "enhanced": enhanced_data,
            "changes": changes,
            "success": True
        }
    except Exception as e:
        print(f"⚠️ Enhancement failed: {str(e)}")
        return {
            "original": resume_data,
            "enhanced": resume_data,
            "changes": ["Enhancement failed — using original data"],
            "success": False
        }

# ==================== URL FETCHING ====================

def fetch_job_description_from_url(url: str) -> str:
    if not SCRAPING_AVAILABLE:
        raise Exception("BeautifulSoup not installed")
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        for element in soup(['script', 'style', 'nav', 'header', 'footer']):
            element.decompose()
        text = soup.get_text(separator='\n', strip=True)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        text = '\n'.join(lines)
        print(f"✅ URL fetched: {len(text)} chars")
        return text
    except Exception as e:
        raise Exception(f"URL fetch failed: {str(e)}")

# ==================== LATEX GENERATION ====================

def generate_latex_from_template(resume_data: dict) -> str:
    def escape_latex(text):
        if not text:
            return ""
        replacements = {
            '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#',
            '_': r'\_', '{': r'\{', '}': r'\}',
            '~': r'\textasciitilde{}', '^': r'\^{}'
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text

    latex = r"""\documentclass{resume}
\usepackage[paperheight=12in,paperwidth=9in,left=0.4in,top=0.4in,right=0.4in,bottom=0.4in]{geometry}
\newcommand{\tab}[1]{\hspace{.2667\textwidth}\rlap{#1}}
\newcommand{\itab}[1]{\hspace{0em}\rlap{#1}}
"""
    latex += f"\\name{{{escape_latex(resume_data.get('name', 'Your Name'))}}}\n"

    phone = resume_data.get('phone', '')
    location = resume_data.get('location', '')
    if phone or location:
        latex += f"\\address{{{escape_latex(phone)} \\\\ {escape_latex(location)}}}\n"

    email = resume_data.get('email', '')
    linkedin = resume_data.get('linkedin', '')
    github = resume_data.get('github', '')
    contact_line = ""
    if email:
        contact_line += f"\\href{{{email}}}{{{escape_latex(email)}}}"
    if linkedin:
        if contact_line: contact_line += " \\\\ "
        contact_line += f"\\href{{{linkedin}}}{{LinkedIn}}"
    if github:
        if contact_line: contact_line += " \\\\ "
        contact_line += f"\\href{{{github}}}{{Github}}"
    if contact_line:
        latex += f"\\address{{{contact_line}}}\n"

    latex += "\\begin{document}\n\n"
    latex += "\\begin{rSection}{Education}\n"
    for edu in resume_data.get('education', []):
        degree = escape_latex(edu.get('degree', ''))
        institution = escape_latex(edu.get('institution', ''))
        year = escape_latex(edu.get('year', ''))
        latex += f"{{\\bf {degree}}}, {institution} \\hfill {{{year}}}\n"
    latex += "\\end{rSection}\n\n"

    latex += "\\begin{rSection}{TECHNICAL SKILLS}\n"
    latex += "\\begin{tabular}{ @{} >{\\bfseries}l @{\\hspace{6ex}} l }\n"
    skills = resume_data.get('technical_skills', {})
    if skills.get('languages'):
        latex += f"LANGUAGES: & {', '.join([escape_latex(s) for s in skills['languages']])}\\\\\n"
    if skills.get('databases'):
        latex += f"DATABASES: & {', '.join([escape_latex(s) for s in skills['databases']])}\\\\\n"
    if skills.get('cloud'):
        latex += f"CLOUD: & {', '.join([escape_latex(s) for s in skills['cloud']])}\\\\\n"
    if skills.get('tools'):
        latex += f"TOOLS: & {', '.join([escape_latex(s) for s in skills['tools']])}\\\\\n"
    latex += "\\end{tabular}\\\\\n\\end{rSection}\n\n"

    latex += "\\begin{rSection}{WORK EXPERIENCE}\n"
    for exp in resume_data.get('work_experience', []):
        title = escape_latex(exp.get('title', ''))
        company = escape_latex(exp.get('company', ''))
        location = escape_latex(exp.get('location', ''))
        duration = escape_latex(exp.get('duration', ''))
        latex += f"\\textbf{{{title}}} \\hfill {duration}\\\\\n"
        latex += f"{company} \\hfill \\textit{{{location}}}\n"
        latex += "\\begin{itemize}\n\\itemsep -3pt {}\n"
        for achievement in exp.get('achievements', []):
            latex += f"\\item {escape_latex(achievement)}\n"
        latex += "\\end{itemize}\n\n"
    latex += "\\end{rSection}\n\n"

    if resume_data.get('projects'):
        latex += "\\begin{rSection}{PERSONAL PROJECTS}\n\\begin{itemize}\n\\itemsep -3pt {}\n"
        for project in resume_data['projects']:
            name = escape_latex(project.get('name', ''))
            desc = escape_latex(project.get('description', ''))
            latex += f"\\item \\textbf{{{name}}} - {desc}\n"
        latex += "\\end{itemize}\n\\end{rSection}\n\n"

    if resume_data.get('achievements'):
        latex += "\\begin{rSection}{ACHIEVEMENTS}\n\\begin{itemize}\n"
        for achievement in resume_data['achievements']:
            latex += f"\\item {escape_latex(achievement)}\n"
        latex += "\\end{itemize}\n\\end{rSection}\n\n"

    if resume_data.get('certifications'):
        latex += "\\begin{rSection}{CERTIFICATION}\n\\begin{itemize}\n"
        for cert in resume_data['certifications']:
            latex += f"\\item {escape_latex(cert)}\n"
        latex += "\\end{itemize}\n\\end{rSection}\n\n"

    latex += "\\end{document}\n"
    return latex

# ==================== ENDPOINTS ====================

@app.post("/parse-resume")
async def parse_resume_endpoint(
    resume: Optional[UploadFile] = File(None)
):
    """Quickly parse resume into structured data for the freshness check modal"""
    try:
        if not resume:
            return {"success": False, "error": "No resume provided"}
        content = await resume.read()
        resume_text = extract_text_from_file(resume.filename, content)
        parsed = parse_resume_to_structured_data(resume_text)
        return {"success": True, "data": parsed}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/analyze")
async def analyze_resume(
    job_description: Optional[UploadFile] = File(None),
    resume: Optional[UploadFile] = File(None),
    jd_text: Optional[str] = Form(None),
    jd_url: Optional[str] = Form(None),
    resume_json: Optional[str] = Form(None),   # Updated resume from frontend wizard
):
    """Complete AI-powered resume analysis"""
    try:
        # Extract JD
        jd_content = ""
        if job_description:
            content = await job_description.read()
            jd_content = extract_text_from_file(job_description.filename, content)
        elif jd_text:
            jd_content = jd_text
        elif jd_url:
            jd_content = fetch_job_description_from_url(jd_url)

        # Extract Resume — prefer wizard-updated JSON if provided
        resume_content = ""
        resume_data_override = None

        if resume_json:
            try:
                resume_data_override = json.loads(resume_json)
                # Convert structured data back to text for AI analysis
                exp_text = "\n".join([
                    f"{e.get('title')} at {e.get('company')} ({e.get('duration')}): " +
                    " ".join(e.get('achievements', []))
                    for e in resume_data_override.get('work_experience', [])
                ])
                skills = resume_data_override.get('technical_skills', {})
                all_skills = ", ".join([s for cat in skills.values() for s in (cat if isinstance(cat, list) else [])])
                resume_content = f"""
Name: {resume_data_override.get('name', '')}
Skills: {all_skills}
Experience:
{exp_text}
Projects: {", ".join([p.get('name', '') for p in resume_data_override.get('projects', [])])}
"""
                print(f"✅ Using wizard-updated resume for {resume_data_override.get('name')}")
            except Exception as e:
                print(f"⚠️ Could not parse resume_json: {e}")

        if not resume_content and resume:
            content = await resume.read()
            resume_content = extract_text_from_file(resume.filename, content)

        if not jd_content or not resume_content:
            return {"success": False, "error": "Both job description and resume are required"}

        # ── Parse resume structure
        print("📊 Parsing resume...")
        resume_data = resume_data_override or parse_resume_to_structured_data(resume_content)

        # ── ONE combined AI call: ATS + gaps + rewrites + skills + summary + interview prep
        print("🤖 Running full analysis (single API call)...")
        full_result = run_full_analysis(jd_content, resume_content)
        full_text = full_result.get("full_analysis", "")

        # ── Parse sections by delimiter
        def extract_section(text, marker, next_markers):
            start = text.find(f"###{marker}###")
            if start == -1:
                return ""
            start = text.find("\n", start) + 1
            end = len(text)
            for nm in next_markers:
                idx = text.find(f"###{nm}###", start)
                if idx != -1 and idx < end:
                    end = idx
            return text[start:end].strip()

        all_markers = ["ATS_SCORE", "CRITICAL_GAPS", "REWRITES", "SKILLS", "SUMMARY", "INTERVIEW_PREP"]
        ats_text       = extract_section(full_text, "ATS_SCORE",      all_markers[1:])
        gaps_text      = extract_section(full_text, "CRITICAL_GAPS",  all_markers[2:])
        rewrites_text  = extract_section(full_text, "REWRITES",       all_markers[3:])
        skills_text    = extract_section(full_text, "SKILLS",         all_markers[4:])
        summary_text   = extract_section(full_text, "SUMMARY",        all_markers[5:])
        interview_text = extract_section(full_text, "INTERVIEW_PREP", [])

        # ── YouTube (lightweight, skip on re-runs)
        if resume_data_override:
            youtube_resources = []
        else:
            print("🎥 Resources...")
            youtube_resources = generate_youtube_resources(jd_content)

        # No more enhancement call — user controls their own rewrites via accept/reject

        return {
            "success": True,
            "data": {
                "ats_score":      {"score_analysis": ats_text},
                "critical_gaps":  gaps_text,
                "rewrites":       rewrites_text,
                "skills":         skills_text,
                "summary":        summary_text,
                "interview_prep": {"interview_prep": interview_text},
                "youtube_resources": youtube_resources,
                "resume_data":    resume_data_override or resume_data,
                "original_resume_data": resume_data
            },
            "message": "Analysis complete!"
        }
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        return {"success": False, "error": str(e)}


@app.post("/download-resume")
async def download_optimized_resume(resume_data: str = Form(...)):
    """Generate LaTeX resume file"""
    try:
        data = json.loads(resume_data)
        latex_content = generate_latex_from_template(data)
        temp_dir = tempfile.mkdtemp()
        tex_path = os.path.join(temp_dir, "optimized_resume.tex")
        with open(tex_path, 'w', encoding='utf-8') as f:
            f.write(latex_content)
        return FileResponse(
            tex_path,
            media_type='application/x-tex',
            filename="optimized_resume.tex"
        )
    except Exception as e:
        return {"success": False, "error": str(e)}